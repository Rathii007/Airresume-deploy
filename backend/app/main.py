from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends, Request, Response
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi_limiter import FastAPILimiter
from pydantic import BaseModel, Field, EmailStr
from starlette.requests import Request
from starlette.responses import Response
from starlette.middleware.base import BaseHTTPMiddleware
import spacy
import PyPDF2
import pytesseract
from PIL import Image
import redis.asyncio as redis
import openai
import os
import logging
import time
import io
import asyncio
import aiohttp
import sqlite3
from datetime import datetime
from io import BytesIO
from dotenv import load_dotenv
from typing import Optional
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from pathlib import Path
from pdf2image import convert_from_bytes
from docx import Document
import uvicorn
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Load environment variables
load_dotenv(dotenv_path=Path(__file__).parent.parent / ".env")

# Configuration
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
REDIS_URL = os.getenv("REDIS_URL", "redis://localhost:6379")
SPACY_MODEL = os.getenv("SPACY_MODEL", "en_core_web_sm")
ALLOWED_ORIGINS = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000").split(",")
ENVIRONMENT = os.getenv("ENVIRONMENT", "development")

if not GROQ_API_KEY:
    raise ValueError("⚠️ GROQ_API_KEY is missing! Set it in your .env file.")

# Configure OpenAI API for Groq
openai.api_key = GROQ_API_KEY
openai.api_base = "https://api.groq.com/openai/v1"

# Initialize FastAPI
app = FastAPI(debug=ENVIRONMENT != "production")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST"],
    allow_headers=["*"],
)

# Logging Configuration
logging.basicConfig(
    level=logging.INFO if ENVIRONMENT == "production" else logging.DEBUG,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger(__name__)

# Middleware for File Upload Size Limit (5 MB)
class LimitUploadSize(BaseHTTPMiddleware):
    def __init__(self, app, max_size: int):
        super().__init__(app)
        self.max_size = max_size

    async def dispatch(self, request: Request, call_next):
        if request.method == "POST":
            content_length = int(request.headers.get("content-length", 0))
            if content_length > self.max_size:
                raise HTTPException(status_code=413, detail="File too large. Maximum size is 5 MB.")
        return await call_next(request)

app.add_middleware(LimitUploadSize, max_size=5 * 1024 * 1024)

# Middleware for Request Timeout (30 seconds)
class TimeoutMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        try:
            return await asyncio.wait_for(call_next(request), timeout=30)
        except asyncio.TimeoutError:
            raise HTTPException(status_code=504, detail="Request timed out")

app.add_middleware(TimeoutMiddleware)

# Middleware to enforce HTTPS by redirecting HTTP to HTTPS
class HTTPSRedirectMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        # Check if the request is over HTTP and the environment is production
        if request.url.scheme == "http" and ENVIRONMENT == "production":
            # Construct the HTTPS URL
            https_url = request.url.replace(scheme="https")
            # Return a redirect response
            return Response(
                status_code=301,
                headers={"Location": str(https_url)},
            )
        return await call_next(request)

app.add_middleware(HTTPSRedirectMiddleware)

# Dependency Injection
async def get_redis(request: Request):
    client = await redis.from_url(REDIS_URL, encoding="utf-8", decode_responses=True)
    try:
        await client.ping()  # Test connection
        yield client
    except redis.ConnectionError as e:
        logger.error("Failed to connect to Redis: %s", str(e))
        yield None
    finally:
        await client.close()

async def get_nlp():
    nlp = spacy.load(SPACY_MODEL)
    yield nlp

# Rate Limiter with IP-based limiting
def RateLimiter(times: int, seconds: int):
    async def limiter(request: Request, redis_client=Depends(get_redis)):
        if redis_client is None:
            logger.warning("Rate limiting not enforced due to Redis unavailability.")
            return
        client_ip = request.client.host
        key = f"ratelimit:{client_ip}:{request.url.path}"
        count = await redis_client.get(key)
        if count is None:
            await redis_client.setex(key, seconds, 1)
        elif int(count) >= times:
            raise HTTPException(status_code=429, detail="Too many requests")
        else:
            await redis_client.incr(key)
    return limiter

# Startup and Shutdown Events
@app.on_event("startup")
async def startup_event():
    # Create SQLite database and table for feedback
    conn = sqlite3.connect("feedback.db")
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            email TEXT,
            feedback TEXT NOT NULL,
            submitted_at TIMESTAMP
        )
    """)
    conn.commit()
    conn.close()

    # Redis startup for rate limiting
    redis_client = await redis.from_url(REDIS_URL, encoding="utf-8", decode_responses=True)
    try:
        await redis_client.ping()
        await FastAPILimiter.init(redis_client)
        logger.info("Application started with Redis connection established.")
    except redis.ConnectionError as e:
        logger.error("Redis connection failed: %s. Shutting down.", str(e))
        raise RuntimeError("Redis is required for rate limiting. Shutting down.")

@app.on_event("shutdown")
async def shutdown_event():
    await FastAPILimiter.close()
    logger.info("Application shutdown, Redis connection closed.")

# Helper Functions
async def ai_suggestions(prompt_text):
    max_retries = 3
    retry_delay = 15
    for attempt in range(max_retries):
        try:
            logger.debug("Calling Groq API with prompt: %s", prompt_text[:50] + "...")
            # Use aiohttp for async HTTP requests with a timeout
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    "https://api.groq.com/openai/v1/chat/completions",
                    headers={"Authorization": f"Bearer {GROQ_API_KEY}"},
                    json={
                        "model": "llama3-8b-8192",
                        "messages": [
                            {"role": "system", "content": "You are an expert resume assistant."},
                            {"role": "user", "content": prompt_text},
                        ]
                    },
                    timeout=10  # 10-second timeout
                ) as response:
                    response_data = await response.json()
                    return response_data["choices"][0]["message"]["content"].strip()
        except aiohttp.ClientTimeout:
            logger.error("Groq API request timed out")
            return "Request to AI service timed out."
        except openai.error.RateLimitError as e:
            if attempt < max_retries - 1:
                logger.warning("Rate limit hit, retrying in %d seconds...", retry_delay)
                await asyncio.sleep(retry_delay)
            else:
                logger.error("Rate limit exceeded after %d attempts.", max_retries)
                return "Rate limit exceeded."
        except Exception as e:
            logger.error("AI Suggestion Error: %s", str(e))
            return "Error generating suggestion."

def extract_text_from_pdf(file_bytes):
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
            else:
                images = convert_from_bytes(file_bytes, first_page=page_num + 1, last_page=page_num + 1)
                if images:
                    text += pytesseract.image_to_string(images[0]) + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {str(e)}")
        return ""

def analyze_structure(text, nlp=Depends(get_nlp)):
    doc = nlp(text)
    sections = {"experience": False, "education": False, "skills": False}
    for token in doc:
        if token.text.lower() in sections:
            sections[token.text.lower()] = True
    return [k for k, v in sections.items() if not v]

def readability_metrics(text, nlp=Depends(get_nlp)):
    doc = nlp(text)
    sentences = list(doc.sents)
    avg_sentence_length = sum(len(sent) for sent in sentences) / len(sentences) if sentences else 0
    return {"avg_sentence_length": avg_sentence_length}

def calculate_ats_score(resume_text, nlp=Depends(get_nlp)):
    doc = nlp(resume_text)
    common_keywords = {"python", "java", "sql", "team", "project", "management", "skills", "experience", "education", "certified"}
    resume_words = set([token.text.lower() for token in doc if token.is_alpha])
    keyword_hits = len(common_keywords.intersection(resume_words))
    keyword_score = min((keyword_hits / 10) * 100, 100)
    missing_sections = analyze_structure(resume_text, nlp)
    structure_score = 100 - (len(missing_sections) * 33)
    metrics = readability_metrics(resume_text, nlp)
    readability_score = 100 if metrics['avg_sentence_length'] < 20 else max(0, 100 - (metrics['avg_sentence_length'] - 20) * 5)
    word_count = len(resume_words)
    length_score = 100 if 150 <= word_count <= 500 else 50 if 100 <= word_count < 150 or 500 < word_count <= 700 else 25
    return round((0.3 * keyword_score) + (0.3 * structure_score) + (0.2 * readability_score) + (0.2 * length_score), 2)

def sanitize_text(text):
    if not isinstance(text, str):
        return ""
    return ''.join(c for c in text if ord(c) >= 32 or c in '\n\r\t')

# Pydantic Models with Stricter Validation
class SuggestionRequest(BaseModel):
    jobTitle: str = Field(..., min_length=2, max_length=100)
    yearsExperience: str = Field(..., pattern=r"^\d+$", max_length=2)

class EnhanceRequest(BaseModel):
    section: str = Field(..., min_length=2, max_length=50)
    content: str = Field(..., min_length=10, max_length=2000)
    jobTitle: Optional[str] = Field(None, min_length=2, max_length=100)

class ResumeData(BaseModel):
    name: str = Field(..., min_length=2, max_length=100)
    email: EmailStr
    phone: str = Field(..., max_length=20)
    education: str = Field(..., max_length=1000)
    experience: str = Field(..., max_length=2000)
    skills: str = Field(..., max_length=1000)
    template: str = Field(..., min_length=2, max_length=50)
    format: Optional[str] = "pdf"

class FeedbackRequest(BaseModel):
    name: Optional[str] = Field(None, max_length=100)
    email: Optional[EmailStr] = None
    feedback: str = Field(..., min_length=10, max_length=1000)

# Improved Resume Extraction
def improved_extract_resume(file_bytes, nlp=Depends(get_nlp)):
    text = extract_text_from_pdf(file_bytes)
    if not text:
        return {"name": "", "email": "", "phone": "", "education": "", "experience": "", "skills": ""}
    
    doc = nlp(text)
    data = {"name": "", "email": "", "phone": "", "education": "", "experience": "", "skills": ""}
    lines = text.split("\n")
    for i, line in enumerate(lines):
        line = line.strip()
        if not data["name"] and line and "@" not in line and not any(c.isdigit() for c in line):
            data["name"] = line
        if "@" in line and not data["email"]:
            data["email"] = line
        if any(c.isdigit() for c in line) and 8 <= len(line.replace(" ", "").replace("-", "")) <= 15 and not data["phone"]:
            data["phone"] = line
        if "education" in line.lower():
            data["education"] = "\n".join(lines[i+1:i+4]).strip() if i+1 < len(lines) else line
        if "experience" in line.lower() or "work" in line.lower():
            data["experience"] = "\n".join(lines[i+1:i+6]).strip() if i+1 < len(lines) else line
        if "skills" in line.lower():
            data["skills"] = "\n".join(lines[i+1:i+4]).strip() if i+1 < len(lines) else line
    return {k: sanitize_text(v) for k, v in data.items()}

# Endpoints
@app.post("/parse-resume/")
async def parse_resume(file: UploadFile, rate_limit=Depends(RateLimiter(times=5, seconds=60))):
    try:
        content = await file.read()
        if file.filename.endswith(".pdf"):
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = "".join(page.extract_text() + "\n" for page in pdf_reader.pages)
        elif file.filename.endswith(".docx"):
            doc = Document(BytesIO(content))
            text = "".join(para.text + "\n" for para in doc.paragraphs)
        else:
            raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")
        
        sections = []
        current_section = {"title": "", "content": ""}
        for line in text.split("\n"):
            if line.isupper() or (len(line) < 30 and line.strip().isalpha()):
                if current_section["title"]:
                    sections.append(current_section)
                current_section = {"title": line.strip(), "content": ""}
            else:
                current_section["content"] += line + "\n"
        if current_section["title"]:
            sections.append(current_section)
        return {"sections": sections}
    except Exception as e:
        logger.error("Parse resume error: %s", str(e))
        raise HTTPException(status_code=500, detail=f"Error parsing file: {str(e)}")

@app.post("/extract-resume/")
async def extract_resume(resume: UploadFile = File(...), nlp=Depends(get_nlp), rate_limit=Depends(RateLimiter(times=5, seconds=60))):
    resume_bytes = await resume.read()
    return improved_extract_resume(resume_bytes, nlp)

@app.post("/suggest-content/")
async def suggest_content(data: SuggestionRequest, rate_limit=Depends(RateLimiter(times=3, seconds=60))):
    prompt = (
        f"Generate a resume for a {data.jobTitle} with {data.yearsExperience} years of experience. "
        "Return in this format:\n"
        "- Education: [One concise entry]\n"
        "- Experience: [3 bullet points]\n"
        "- Skills: [5 concise skills]"
    )
    response = await ai_suggestions(prompt)
    data = {"education": "", "experience": "", "skills": ""}
    for line in response.split("\n"):
        if line.startswith("- Education:"):
            data["education"] = sanitize_text(line.replace("- Education:", "").strip())
        elif line.startswith("- Experience:"):
            data["experience"] = sanitize_text("\n".join([l.strip() for l in response.split("\n") if l.startswith("- ") and "Experience" not in l][:3]))
        elif line.startswith("- Skills:"):
            data["skills"] = sanitize_text("\n".join([l.strip() for l in response.split("\n") if l.startswith("- ") and "Skills" not in l][:5]))
    return data

@app.post("/enhance-section/")
async def enhance_section(data: EnhanceRequest, rate_limit=Depends(RateLimiter(times=5, seconds=60))):
    context = f" for a {data.jobTitle}" if data.jobTitle else ""
    prompt = (
        f"Rewrite this {data.section} section{context} to be concise, professional, and ATS-friendly:\n"
        f"{data.content}\n"
        f"Return the enhanced version only, no extra text."
    )
    enhanced = await ai_suggestions(prompt)
    return {"enhanced": sanitize_text(enhanced)}

@app.post("/ats-preview/")
async def ats_preview(data: ResumeData, nlp=Depends(get_nlp), rate_limit=Depends(RateLimiter(times=5, seconds=60))):
    resume_text = f"{data.name}\n{data.email}\n{data.phone}\n{data.education}\n{data.experience}\n{data.skills}"
    ats_score = calculate_ats_score(resume_text, nlp)
    return {"ats_score": ats_score}

@app.post("/generate-resume/")
async def generate_resume(data: ResumeData, rate_limit=Depends(RateLimiter(times=10, seconds=60))):
    try:
        data_dict = data.dict()
        format = data_dict.pop("format", "pdf")
        
        # Sanitize input data
        sanitized_data = {k: sanitize_text(v) for k, v in data_dict.items()}
        
        # Define available templates
        template_styles = {
            "modern": {"font": "Helvetica", "color": "blue"},
            "classic": {"font": "Times", "color": "black"},
            "creative": {"font": "Helvetica", "color": "purple"},
            "executive": {"font": "Times", "color": "darkgray"},
            "minimalist": {"font": "Helvetica", "color": "black"},
        }
        
        if sanitized_data.get("template") not in template_styles:
            raise HTTPException(
                status_code=400,
                detail=f"Template '{sanitized_data.get('template')}' not found. Available templates: {list(template_styles.keys())}"
            )

        if format == "pdf":
            # Generate PDF in memory using reportlab
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            
            # Apply basic styling based on template
            c.setFont(template_styles[sanitized_data["template"]]["font"], 12)
            c.setFillColor(template_styles[sanitized_data["template"]]["color"])
            
            # Add content
            y_position = height - 50
            line_height = 15
            
            # Name
            c.setFontSize(16)
            c.drawString(50, y_position, sanitized_data.get("name", "Your Name"))
            y_position -= line_height * 2
            
            # Contact Info
            c.setFontSize(12)
            c.drawString(50, y_position, f"{sanitized_data.get('email', '')} | {sanitized_data.get('phone', '')}")
            y_position -= line_height * 2
            
            # Education
            c.setFontSize(14)
            c.drawString(50, y_position, "Education")
            y_position -= line_height
            c.setFontSize(12)
            for line in sanitized_data.get("education", "Add your education").split("\n"):
                c.drawString(50, y_position, line)
                y_position -= line_height
            y_position -= line_height
            
            # Experience
            c.setFontSize(14)
            c.drawString(50, y_position, "Experience")
            y_position -= line_height
            c.setFontSize(12)
            for line in sanitized_data.get("experience", "Add your experience").split("\n"):
                c.drawString(50, y_position, line)
                y_position -= line_height
            y_position -= line_height
            
            # Skills
            c.setFontSize(14)
            c.drawString(50, y_position, "Skills")
            y_position -= line_height
            c.setFontSize(12)
            for line in sanitized_data.get("skills", "Add your skills").split("\n"):
                c.drawString(50, y_position, line)
                y_position -= line_height
            
            c.showPage()
            c.save()
            
            buffer.seek(0)
            return StreamingResponse(
                buffer,
                media_type="application/pdf",
                headers={"Content-Disposition": "attachment; filename=resume.pdf"}
            )
        
        elif format == "docx":
            # Generate DOCX in memory using python-docx
            doc = Document()
            
            # Add content
            doc.add_heading(sanitized_data.get("name", "Your Name"), level=1)
            doc.add_paragraph(f"{sanitized_data.get('email', '')} | {sanitized_data.get('phone', '')}")
            doc.add_heading("Education", level=2)
            doc.add_paragraph(sanitized_data.get("education", "Add your education"))
            doc.add_heading("Experience", level=2)
            doc.add_paragraph(sanitized_data.get("experience", "Add your experience"))
            doc.add_heading("Skills", level=2)
            doc.add_paragraph(sanitized_data.get("skills", "Add your skills"))
            
            # Save to BytesIO buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return StreamingResponse(
                buffer,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=resume.docx"}
            )
        
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")
    
    except Exception as e:
        logger.error(f"Generate resume failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate resume: {str(e)}")

@app.post("/match-resume/")
async def match_resume(
    resume: UploadFile = File(...),
    job_description: Optional[str] = Form(default=None, max_length=5000),
    nlp=Depends(get_nlp),
    rate_limit=Depends(RateLimiter(times=3, seconds=60))
):
    try:
        resume_content = await resume.read()
        resume_text = extract_text_from_pdf(resume_content)
        if not resume_text:
            raise HTTPException(status_code=400, detail="No text found in the resume PDF")

        missing_sections = analyze_structure(resume_text, nlp)
        structure_score = max(0, 100 - len(missing_sections) * 20)
        structure_prompt = (
            f"Analyze the structure of this resume:\n{resume_text}\n"
            "Comment on missing sections like Experience, Education, or Skills in 1-2 concise sentences. "
            "Then, provide 5-7 actionable suggestions as bullet points starting with '- '."
        ) if job_description else (
            f"Analyze the structure of this resume:\n{resume_text}\n"
            "Comment on missing sections like Experience, Education, or Skills in 2-3 concise sentences. "
            "Then, provide 3-5 actionable suggestions as bullet points starting with '- '."
        )
        structure_feedback = await ai_suggestions(structure_prompt)

        metrics = readability_metrics(resume_text, nlp)
        readability_score = min(100, max(0, 100 - (metrics["avg_sentence_length"] - 15) * 5))
        readability_prompt = (
            f"Analyze the readability of this resume:\n{resume_text}\n"
            f"Average sentence length is {metrics['avg_sentence_length']:.1f} words. "
            "Provide feedback on clarity and sentence structure in 1-2 concise sentences and every new suggestion should starts from a new line. "
            "Then, provide 5-7 actionable suggestions as bullet points starting with '- ' and every new suggestion should starts from a new line."
        ) if job_description else (
            f"Analyze the readability of this resume:\n{resume_text}\n"
            f"Average sentence length is {metrics['avg_sentence_length']:.1f} words. "
            "Provide feedback on clarity and sentence structure in 2-3 concise sentences and every new suggestion should starts from a new line. "
            "Then, provide 3-5 actionable suggestions as bullet points starting with '- ' and every new suggestion should starts from a new line."
        )
        readability_feedback = await ai_suggestions(readability_prompt)

        word_count = len(resume_text.split())
        length_score = min(100, max(0, 100 - abs(word_count - 500) * 0.2))

        if job_description:
            def preprocess(text):
                doc = nlp(text.lower())
                return " ".join([token.lemma_ for token in doc if not token.is_stop and token.is_alpha])

            processed_resume = preprocess(resume_text)
            processed_job_desc = preprocess(job_description)
            vectorizer = TfidfVectorizer()
            tfidf_matrix = vectorizer.fit_transform([processed_resume, processed_job_desc])
            match_score = cosine_similarity(tfidf_matrix[0], tfidf_matrix[1])[0][0] * 100
            job_keywords = set(processed_job_desc.split())
            resume_keywords = set(processed_resume.split())
            missing_keywords = list(job_keywords - resume_keywords)
            keyword_score = min(100, len(resume_keywords.intersection(job_keywords)) * 10)

            overall_match_score = int(
                (keyword_score * 0.4) + (structure_score * 0.3) + (readability_score * 0.2) + (length_score * 0.1)
            )

            ai_prompt = (
                f"You are an expert resume reviewer. Compare this resume:\n{resume_text}\n"
                f"with this job description:\n{job_description}. "
                "Provide feedback in the following format:\n"
                "**Match Quality and Suggestions for Improvement:**\n[Explain the match quality in 1-2 concise sentences. Then, provide 5-7 actionable suggestions as bullet points starting with '- ' and after completion start the next suggestion in a new line.]\n"
                "**Overall Quality, Clarity, and Structure:**\n[Analyze the resume’s overall quality, clarity, and structure in 1-2 concise sentences. After every suggestion, start the next suggestion in a new line.]"
            )
            ai_feedback = await ai_suggestions(ai_prompt)

            strengths = "No strengths identified."
            overall_quality = "No overall quality feedback provided."
            if "**Match Quality and Suggestions for Improvement:**" in ai_feedback:
                feedback_parts = ai_feedback.split("**Match Quality and Suggestions for Improvement:**")
                if len(feedback_parts) > 1:
                    strengths_part = feedback_parts[1]
                    if "**Overall Quality, Clarity, and Structure:**" in strengths_part:
                        strengths = strengths_part.split("**Overall Quality, Clarity, and Structure:**")[0].strip()
                        overall_quality = ai_feedback.split("**Overall Quality, Clarity, and Structure:**")[1].strip()
                    else:
                        strengths = strengths_part.strip()

            return {
                "match_score": overall_match_score,
                "match_score_raw": f"{match_score:.2f}%",
                "missing_keywords": missing_keywords[:10],
                "explanation": ai_feedback.split('.')[0] + "." if ai_feedback else "No explanation provided.",
                "structure_feedback": structure_feedback,
                "readability_feedback": readability_feedback,
                "metrics": {
                    "keyword_score": keyword_score,
                    "structure_score": structure_score,
                    "readability_score": readability_score,
                    "length_score": length_score,
                    "avg_sentence_length": metrics["avg_sentence_length"],
                },
                "ai_feedback": {
                    "strengths": strengths,
                    "overall_quality": overall_quality,
                }
            }
        else:
            ats_score = calculate_ats_score(resume_text, nlp)
            keyword_score = 50
            overall_ats_score = int(
                (keyword_score * 0.4) + (structure_score * 0.3) + (readability_score * 0.2) + (length_score * 0.1)
            )

            ai_prompt = (
                f"You are an expert resume reviewer. Analyze this resume:\n{resume_text}\n"
                "for ATS compatibility and overall quality. Provide feedback in the following format:\n"
                "**ATS Readiness:**\n[Explain the ATS compatibility in 2-3 concise sentences.]\n"
                "**Suggestions:**\n[Provide 3-5 actionable suggestions to improve structure, keyword usage, and readability as bullet points starting with '- ' and every new suggestion suggestion should start from a new line.]"
            )
            ai_feedback = await ai_suggestions(ai_prompt)

            suggestions = "No specific suggestions provided."
            ats_readiness = "No ATS readiness feedback provided."
            if "**ATS Readiness:**" in ai_feedback:
                ats_parts = ai_feedback.split("**ATS Readiness:**")
                if len(ats_parts) > 1:
                    ats_readiness_part = ats_parts[1]
                    if "**Suggestions:**" in ats_readiness_part:
                        ats_readiness = ats_readiness_part.split("**Suggestions:**")[0].strip()
                        suggestions = ai_feedback.split("**Suggestions:**")[1].strip()
                    else:
                        ats_readiness = ats_readiness_part.strip()

            return {
                "ats_score": overall_ats_score,
                "ats_score_raw": f"{ats_score}%",
                "explanation": ai_feedback.split('.')[0] + "." if ai_feedback else "No explanation provided.",
                "structure_feedback": structure_feedback,
                "readability_feedback": readability_feedback,
                "metrics": {
                    "keyword_score": keyword_score,
                    "structure_score": structure_score,
                    "readability_score": readability_score,
                    "length_score": length_score,
                    "avg_sentence_length": metrics["avg_sentence_length"],
                },
                "ai_feedback": {
                    "ats_readiness": ats_readiness,
                    "suggestions": suggestions,
                }
            }
    except Exception as e:
        logger.error(f"Job matching failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to process resume: {str(e)}")

@app.post("/resume-roast/")
async def resume_roast(
    file: UploadFile = File(...),
    roast_level: Optional[str] = Form(default="spicy", max_length=10),
    nlp=Depends(get_nlp),
    rate_limit=Depends(RateLimiter(times=3, seconds=60))
):
    try:
        resume_text = extract_text_from_pdf(await file.read())
        if not resume_text:
            return {"roast": "A blank resume? Wow, you’re really letting your *nothingness* shine!"}
        missing_sections = analyze_structure(resume_text, nlp)
        metrics = readability_metrics(resume_text, nlp)

        tone_instruction = {
            "mild": "Provide a gentle, lighthearted roast that’s encouraging but still funny.",
            "spicy": "Provide a moderately sarcastic roast that’s humorous but not too harsh.",
            "burnt": "Provide a savage, no-holds-barred roast that’s still funny but very critical."
        }.get(roast_level.lower(), "Provide a moderately sarcastic roast that’s humorous but not too harsh.")

        ai_prompt = (
            f"You’re a stand-up comedian roasting this resume:\n{resume_text}\n"
            f"Missing sections: {', '.join(missing_sections) if missing_sections else 'None'}.\n"
            f"Average sentence length is {metrics['avg_sentence_length']:.1f} words.\n"
            f"{tone_instruction}\n"
            "Provide a concise roast in EXACTLY the following format (use the exact section headers and prefixes):\n"
            "- **Structure**: Roast the layout, formatting, and missing sections in 2-3 sentences.\n"
            "- **Readability**: Roast the clarity, jargon, and verbosity in 2-3 sentences.\n"
            "- **Projects**: Roast the projects section in 2-3 sentences.\n"
            "- **Skills**: Roast the technical skills in 2-3 sentences.\n"
            "- **Overall Vibe**: Summarize the overall impression sarcastically in 2-3 sentences.\n"
            "Ensure each section is short and punchy. Do not deviate from the specified format, even if a section is missing in the resume."
        )
        ai_roast = await ai_suggestions(ai_prompt)
        return {"roast": sanitize_text(ai_roast)}
    except Exception as e:
        logger.error(f"Resume roast failed: {str(e)}")
        return {"roast": "This resume broke me—guess it’s *that* tragic!"}

@app.post("/feedback/")
async def submit_feedback(feedback: FeedbackRequest, rate_limit=Depends(RateLimiter(times=5, seconds=60))):
    try:
        feedback_data = feedback.dict()
        feedback_data["submitted_at"] = datetime.utcnow().isoformat()

        # Log feedback to SQLite
        conn = sqlite3.connect("feedback.db")
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO feedback (name, email, feedback, submitted_at) VALUES (?, ?, ?, ?)",
            (
                feedback_data["name"],
                feedback_data["email"],
                feedback_data["feedback"],
                feedback_data["submitted_at"],
            ),
        )
        conn.commit()
        conn.close()

        logger.info(f"Feedback received: {feedback_data}")
        return {"message": "Thank you for your feedback!"}
    except Exception as e:
        logger.error(f"Feedback submission failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to submit feedback")

@app.get("/")
async def root():
    return {"message": "FastAPI Resume API is running!"}

# Run the app
if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000, log_level="info" if ENVIRONMENT == "production" else "debug")